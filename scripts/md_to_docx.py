"""One-off: convert golden_brd.md to Bedrock/GoldenBRD.docx."""
from pathlib import Path
from docx import Document
from docx.shared import Pt

def md_to_docx(md_path: Path, docx_path: Path) -> None:
  text = md_path.read_text(encoding="utf-8")
  doc = Document()
  for block in text.split("\n\n"):
    block = block.strip()
    if not block:
      continue
    lines = block.split("\n")
    first = lines[0]
    if first.startswith("# "):
      doc.add_heading(first[2:].strip(), level=0)
      for line in lines[1:]:
        if line.startswith("## "):
          doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
          doc.add_heading(line[4:].strip(), level=2)
        else:
          doc.add_paragraph(line)
    elif first.startswith("## "):
      doc.add_heading(first[3:].strip(), level=1)
      for line in lines[1:]:
        if line.startswith("### "):
          doc.add_heading(line[4:].strip(), level=2)
        else:
          doc.add_paragraph(line)
    elif first.startswith("### "):
      doc.add_heading(first[4:].strip(), level=2)
      for line in lines[1:]:
        doc.add_paragraph(line)
    else:
      doc.add_paragraph(block.replace("\n", " "))
  doc.save(docx_path)

if __name__ == "__main__":
  root = Path(__file__).resolve().parents[1]
  md_path = root / "example_data" / "golden_brd.md"
  docx_path = root / "Bedrock" / "GoldenBRD.docx"
  docx_path.parent.mkdir(parents=True, exist_ok=True)
  md_to_docx(md_path, docx_path)
  print("Wrote", docx_path)
